# -*- coding: utf-8 -*-
"""Document inventory and run-level text extraction engine."""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Set

from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class Location:
    """Base class for a text location inside docx."""
    pass


@dataclass
class ParaLocation(Location):
    """Location inside a paragraph."""
    para_idx: int

    def __repr__(self) -> str:
        return f"Para[{self.para_idx}]"


@dataclass
class CellLocation(Location):
    """Location inside a table cell."""
    table_idx: int
    row_idx: int
    cell_idx: int

    def __repr__(self) -> str:
        return f"Table[{self.table_idx}].Row[{self.row_idx}].Cell[{self.cell_idx}]"


@dataclass
class RunInfo:
    """Describes a single run within a paragraph for formatting preservation."""
    run_index: int
    char_offset: int
    char_length: int


@dataclass
class TextBlock:
    """Extracted text block with back-references to docx DOM."""
    text: str
    location: Location
    runs: List[RunInfo] = field(default_factory=list)
    paragraph_ref: Optional[Paragraph] = field(default=None, repr=False)


@dataclass
class PIISpan:
    """Detected PII span in a TextBlock."""
    start: int
    end: int
    pii_type: str
    matched_text: str
    block_index: int

    def __repr__(self) -> str:
        return f"PIISpan({self.pii_type!r}, {self.start}:{self.end}, {self.matched_text!r})"

    def to_dict(self) -> Dict[str, Any]:
        """Serialize span for JSON metadata logging."""
        return {
            "start": self.start,
            "end": self.end,
            "pii_type": self.pii_type,
            "matched_text": self.matched_text,
            "block_index": self.block_index,
        }


class TextInventory:
    """Walks the docx document, extracting paragraphs & table cells into TextBlocks."""

    def __init__(self, doc: Document) -> None:
        self.doc = doc
        self.blocks: List[TextBlock] = []
        self._extract()

    def _extract(self) -> None:
        """Extract text from all paragraphs and tables."""
        # Paragraphs
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text
            if not text.strip():
                continue
            runs = self._extract_runs(para)
            self.blocks.append(TextBlock(
                text=text,
                location=ParaLocation(para_idx=idx),
                runs=runs,
                paragraph_ref=para,
            ))

        # Tables (including nested tables)
        for t_idx, table in enumerate(self.doc.tables):
            self._extract_table(table, t_idx)

    def _extract_table(self, table: Table, t_idx: int) -> None:
        """Extract text from table cells, avoiding duplicate merged cells."""
        seen_cells: Set[int] = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_id = id(cell)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)

                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text
                    if not text.strip():
                        continue
                    runs = self._extract_runs(para)
                    self.blocks.append(TextBlock(
                        text=text,
                        location=CellLocation(table_idx=t_idx, row_idx=r_idx, cell_idx=c_idx),
                        runs=runs,
                        paragraph_ref=para,
                    ))

                for nested in cell.tables:
                    self._extract_table(nested, t_idx)

    @staticmethod
    def _extract_runs(para: Paragraph) -> List[RunInfo]:
        """Map character offsets to runs for a paragraph."""
        runs: List[RunInfo] = []
        offset = 0
        for ri, run in enumerate(para.runs):
            rlen = len(run.text)
            runs.append(RunInfo(run_index=ri, char_offset=offset, char_length=rlen))
            offset += rlen
        return runs
